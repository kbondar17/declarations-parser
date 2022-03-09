
import requests
import re
import time
import os
import urllib.request
from typing import Union
import csv

import io
import numpy as np
from docx import Document

import pandas as pd

from pathlib import Path
# import tqdm.notebook.tqdm as tqdm
from tqdm import tqdm_notebook


class DataCleaner:
    """убирает лишние данные"""

    @staticmethod
    def remove_unwanted_symbols(df):
        # TODO: чистка всех колонок
        df = df.applymap(lambda x: str(x).replace('\n', ' '))
        return df

    @staticmethod
    def remove_unwanted_cells(df):
        # убирает ячейки с нумерацией
        # print('--- DataCleaner.remove_unwanted_cells ---', df.columns)
        df = df[~df['position'].astype(str).str.isdigit()]
        return df

    def clean_df(self, df):
        df = self.remove_unwanted_symbols(df)
        df = self.remove_unwanted_cells(df)
        return df


class CorrectHeadersParser:

    '''класс для парсинга таблиц, у которых на месте колонки, которые нам нужны'''

    def table_splitter(self, table: pd.DataFrame, file_name) -> list[pd.DataFrame]:
        '''разделяет таблицы, в которых учреждение указано внутри таблицы'''

        def check_if_same(my_array: list) -> bool:
            '''проверяем одинаковые ли колонки'''

            first = my_array[0]
            for e in my_array[1:]:
                if e != first:
                    return False
            return True

        def get_indexes_to_split(table):
            index_to_split = []
            for e in range(len(table)):
                cols = table.iloc[e, :].values
                if check_if_same(cols):
                    index_to_split.append(e)
            return index_to_split

        def split_table(table: pd.DataFrame, index_to_split: Union[int, list[int]], file_name) -> list[pd.DataFrame]:
            dfs = np.array_split(table, index_to_split)
            dfs = [e for e in dfs if len(e) > 0]

            result_dfs = []
            for df in dfs:
                office = df.iloc[0, :][0]
                df = df.iloc[1:, :]
                df['office'] = office
                result_dfs.append(df)

            result_dfs = [e for e in result_dfs if not e.empty]
            try:
                result_dfs = pd.concat(result_dfs)
                return result_dfs
            except Exception as ex:
                print(ex)
                print('rogue file---', file_name)

        index_to_split = get_indexes_to_split(table)

        if not index_to_split:
            return table

        splitted_dfs = split_table(table, index_to_split, file_name)
        return splitted_dfs

    def concat_name(self, df: pd.DataFrame) -> pd.DataFrame:
        '''соединяем колонки ФИО, если они в разных'''

        if 'name' not in df.columns:
            return df

        names_df = df['name']

        check_if_cols_ok = names_df[names_df]

        if isinstance(names_df, str) or isinstance(names_df, pd.Series):
            return df
        # TODO:
        # дропнуть маленькую колонку

        print(type(names_df.values), names_df.values)

        names = [' '.join(e) for e in names_df.values]

        df.drop(columns=['name'], inplace=True)
        df['name'] = names
        return df

    def parse(self, table: pd.DataFrame, file_name) -> pd.DataFrame:
        table = self.concat_name(table)
        table = self.table_splitter(table, file_name)
        return table


class DocxParser:

    def __init__(self):
        self.cols_we_need = ['name', 'salary', 'position', 'department']
        self.parse_correct_headers = CorrectHeadersParser()
        self.parse_incorrect_headers = ''
        self.all_docs: list[dict[str, Document]]
        self.data_cleaner = DataCleaner()

    # def read_all_docs(self, path: str)->list[dict[str, Document]]:
    #     self.path = path
    #     all_docs = []
    #     i = 0
    #     for doc in os.listdir(path)[:50]:
    #         if not doc.endswith('docx'):
    #             continue

    #         try:
    #             all_docs.append({'file':doc,'doc':Document(folder + doc)})
    #             i +=1
    #         except:
    #             print('не вышло--', doc)

    #     self.all_docs = all_docs
    #     print(f'нашли и загрузили {i} файлов')

    def read_docx_tables(self, filename, tab_id=None, **kwargs) -> pd.DataFrame:
        """
        parse table(s) fromt.columnsrd Document (.docx) into Pandas DataFrame(s)

        Parameters:
            filename:   file name of a Word Document

            tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                        When [None] - return a list of DataFrames (parse all tables)

            kwargs:     arguments to pass to `pd.read_csv()` function

        Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
        """
        def read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            return pd.read_csv(vf, **kwargs)

        doc = Document(filename)
        if tab_id is None:
            return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
        else:
            try:
                return read_docx_tab(doc.tables[tab_id], **kwargs)
            except IndexError:
                print(
                    'Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    @staticmethod
    def rename_col(col: str) -> str:

        print('col before rename cols --', col)
        col = col.lower()
        if re.search(pattern='(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество)', string=col):
            return "name"

        elif re.search(pattern='(cреднемесячная|зарпл.|плат[ы, а]|заработн[ой, ая] плат[а, ы]|cреднемесячн[ая, ой]|зарплат[а, ной, ы])', string=col):
            return "salary"

        elif re.search(pattern='(должност[ь, и, ей])', string=col):

            return 'position'

        elif re.search(pattern='(предприяти[е,я]|учреждени[е,я]|юридическое лицо)', string=col):
            return 'department'

        return col

    @staticmethod
    def check_if_columns_ok(cols: tuple) -> bool:
        '''проверяем, есть ли в заголовках таблицы название предприятия и другая инфа'''

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))
        print('колонки на проверке---', cols)
        ok_cols = 0
        company_found = False
        for col in cols:
            company_pattern = '(предприяти[е,я]|учреждени[е,я]|юридическ[ое,ие])'
            res = re.search(pattern=company_pattern, string=col)
            if res:
                company_found = True
                continue

            name_salary_position_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|плат[ы, а]|заработная|плата|cреднемесячн[ая, ой]|зарплат[а, ной, ы]|должность|)'

            res = re.search(pattern=name_salary_position_pattern, string=col)
            if res:
                ok_cols += 1

        if company_found and ok_cols > 1:
            return True
        return False

    def get_all_tables_of_a_doc(self, filename: str) -> list[pd.DataFrame]:

        tables = self.read_docx_tables(filename)
        return tables

    def parse_doc(self, filename: str) -> pd.DataFrame:
        assert filename.endswith('docx'), 'Формат должен быть .docx!'

        doc_tables = self.get_all_tables_of_a_doc(filename)
        parsed_tables = []
        i = 0
        for table in doc_tables:
            i += 1
            # print('table---', table.columns)
            # table.to_excel(f'{i}_промежуточны_странный.xlsx')
            # проверяем норм ли заголовки
            columns_ok = self.check_if_columns_ok(table)
            if not columns_ok:
                # если учреждения нет - смотрим параграфы.
                # добавить документ в опущенные
                pass

            else:
                # если заголовки ок, оставляем только нужные

                table.columns = [self.rename_col(col) for col in table.columns]

                cols_to_leave = [
                    col for col in table.columns if col in self.cols_we_need]
                cols_to_leave = set(cols_to_leave)
                table = table[cols_to_leave]

                # проверяем на наличие вложенных таблиц и фио, разнесенных на несколько стаоблцов
                table = self.parse_correct_headers.parse(table, filename)
                # убираем лишние ячейки и символы
                table = self.data_cleaner.clean_df(table)
                parsed_tables.append(table)

        parsed_tables = [e for e in parsed_tables if isinstance(
            e, pd.DataFrame) and not e.empty]

        if isinstance(parsed_tables, list):
            if parsed_tables:
                concat_tables = pd.concat(parsed_tables)
                return concat_tables

        elif isinstance(parsed_tables, pd.DataFrame):
            if not parsed_tables.empty:
                return concat_tables

    def parse_folder(self, path: str):
        all_tables = []
        for doc in os.listdir(path):
            pass
        # for doc in self.all_docs:
        #     tables = self.parse_doc(doc)
        #     if tables:
        #         # return tables
        #         #print(tables)
        #         # tables = pd.concat(tables)
        #         all_tables.append({'file_name':doc, 'df':tables})

        # return all_tables


# file = "186956_2020_Rukovoditeli_podvedomstvennykh_uchrezhdenii_(kul'tura).docx"
# file = "102907_2019_Rukovoditeli_podvedomstvennykh_uchrezhdenii_(kul'tura).docx"
# file = '101058_2019_Rukovoditeli_podvedomstvennykh_uchrezhdenii_(obrazovanie).docx'
# folder = 'data_ids/docx/'
parser = DocxParser()
# # res = parser.parse_doc(folder + file)
# res = parser.parse_doc(folder + file)
# folder = 'data_ids/doc/'


# res = parser.parse_doc(folder + file)
# res.to_excel(folder + 'cool/ok/' + file + '.xlsx')

folder = 'data_ids/pdf/converted/'
file = '189843_2020_Rukovoditeli.docx'

res = parser.parse_doc(folder + file)

print(res)

fuck = 'me'
