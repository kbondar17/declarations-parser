

import traceback
import sys

from docx2python import docx2python
import requests
import re
import time
import os
import urllib.request
from typing import Union
import csv
import statistics
import pdf2docx

import pickle

import camelot
import typing
import io
import numpy as np
from docx import Document
import matplotlib
import numpy as np
import matplotlib.pyplot as plt

import pandas as pd

from pathlib import Path
# import tqdm.notebook.tqdm as tqdm
from tqdm import tqdm_notebook

import matplotlib.pyplot as plt
import logging
logging.basicConfig(format=u'%(filename)+13s [ LINE:%(lineno)-4s] %(levelname)-8s %(message)s',
                    level=logging.DEBUG)
logger = logging.getLogger(__name__)


class PdfParser:

    @staticmethod
    def convert_pdf_to_df(filename) -> list[pd.DataFrame]:
        tables = camelot.read_pdf(str(filename), line_tol=2, joint_tol=10, line_scale=40, copy_text=[
                                  'v'], pages='1-end')  # , flavor='stream' row_tol=10
        tables = [e.df for e in tables]
        return tables


class CorrectHeadersParser:

    '''класс для парсинга таблиц, у которых на месте колонки, которые нам нужны'''

    def table_splitter(self, table: pd.DataFrame) -> list[pd.DataFrame]:
        '''разделяет таблицы, в которых учреждение указано внутри таблицы'''

        def check_if_same(my_array: list) -> bool:
            '''проверяем одинаковые ли колонки'''

            if len(set(my_array)) > 1:
                return False
            return True

            # first = my_array[0]
            # for e in my_array[1:]:
            #     if e != first:
            #         return False
            # return True

        def get_indexes_to_split(table):
            index_to_split = []
            for e in range(len(table)):
                cols = table.iloc[e, :].values
                if check_if_same(cols):
                    index_to_split.append(e)
            return index_to_split

        def split_table(table: pd.DataFrame, index_to_split: Union[int, list[int]]) -> list[pd.DataFrame]:
            """разделяет таблицу в случае когда название учреждения поместили в середину вот так:

                -должность-  -имя-  -зарплата-
                        -ГБОУ школа 112-
                 директор     Ваня    100 руб

             """
            dfs = np.array_split(table, index_to_split)
            dfs = [e for e in dfs if len(e) > 0]

            result_dfs = []
            for df in dfs:
                office = df.iloc[0, :][0]
                df = df.iloc[1:, :]
                df['office'] = office
                result_dfs.append(df)

            result_dfs = [e for e in result_dfs if not e.empty]
            try:
                result_dfs = pd.concat(result_dfs)
                return result_dfs
            except Exception as ex:
                print(ex)
                print('rogue file---', table)

        index_to_split = get_indexes_to_split(table)

        if not index_to_split:
            return table

        splitted_dfs = split_table(table, index_to_split)
        return splitted_dfs

    def concat_name(self, df: pd.DataFrame) -> pd.DataFrame:
        '''соединяем колонки ФИО, если они в разных'''

        if 'name' not in df.columns:
            return df

        names_df = df['name']

        if isinstance(names_df, str) or isinstance(names_df, pd.Series):
            return df

        # TODO:
        # дропнуть маленькую колонку

        names = [' '.join(e) for e in names_df.values]

        df.drop(columns=['name'], inplace=True)
        df['name'] = names
        return df

    def parse(self, table: pd.DataFrame) -> pd.DataFrame:
        # table = self.concat_name(table)
        table = self.table_splitter(table)
        return table


class DataCleaner:
    """убирает лишние данные"""

    @staticmethod
    def remove_unwanted_symbols(df):
        # TODO: чистка всех колонок
        df = df.applymap(lambda x: str(x).replace('\n', ' '))
        return df

    @staticmethod
    def remove_unwanted_cells(df):
        # убирает ячейки с нумерацией
        # print('--- DataCleaner.remove_unwanted_cells ---', df.columns)
        # TODO: почему тут только должность?
        df = df[~df['position'].astype(str).str.isdigit()]
        return df

    @staticmethod
    def remove_short_rows(df):
        # удаляет ряды с недостаточными данными
        # ! должно применяться после выбора норм колонок
        to_remove = []
        for tup in df.itertuples():
            res = [len(str(e)) for e in tup]
            if statistics.mean(res) < 5:
                to_remove.append(tup.Index)

        df.drop(to_remove, inplace=True)
        return df

    def clean_df(self, df):
        df = self.remove_unwanted_symbols(df)
        df = self.remove_unwanted_cells(df)

        df = self.remove_short_rows(df)
        # print('!!!',df)

        return df


class DocxParser:

    def get_docx_tables(self, filename, tab_id=None, **kwargs) -> list[pd.DataFrame]:
        """
            filename:   file name of a Word Document
            tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                        When [None] - return a list of DataFrames (parse all tables)
        """
        def read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            return pd.read_csv(vf, **kwargs)

        doc = Document(filename)
        if tab_id is None:
            return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
        else:
            try:
                return read_docx_tab(doc.tables[tab_id], **kwargs)
            except IndexError:
                print(
                    'Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    def convert_docx_to_df(self, filename: str) -> pd.DataFrame:
        assert filename.endswith('docx'), 'Формат должен быть .docx!'

        doc = Document(filename)
        # TODO: тут взять текст, который потом прикрутить к

        doc_tables = self.get_docx_tables(filename)

        return doc_tables


class IncorrectHeaders:
    """класс для таблиц с неопределенными заголовками.
        1. Пытаемся найти название учреждений в объединенных ячейках.
        2. Если не получается, для учреждения берем текст, предшествующий таблице. 
    """

    # если прошелся по таблице и нашел вложения внутри - пусть это будет офис.
    # если не нашел - берем название офиса из абзацев вокруг таблиц (если их число плюс минус совпадает)

    def __init__(self):
        pass
        # TODO: добавить обработку чисто docx

        # self.docx_parser = DocxParser()
        # self.pdf_parser = PdfParser()

    def table_splitter(self, table: pd.DataFrame) -> tuple[bool, list[pd.DataFrame]]:
        '''разделяет таблицы, в которых учреждение указано внутри таблицы'''

        def check_if_same(my_array: list) -> bool:
            '''проверяем одинаковые ли колонки'''

            if len(set(my_array)) > 1:
                return False
            return True

        def get_indexes_to_split(table):
            '''определяем индекс строки таблицы, по которому надо разделить'''

            index_to_split = []
            for e in range(len(table)):
                cols = table.iloc[e, :].values
                if check_if_same(cols):
                    index_to_split.append(e)
            return index_to_split

            # если разделили и нашли офис - True
        def split_table(table: pd.DataFrame, index_to_split: Union[int, list[int]]) -> tuple[bool, list[pd.DataFrame]]:
            """разделяет таблицу в случае когда название учреждения поместили в середину вот так:
                -должность-  -имя-  -зарплата-
                        -ГБОУ школа 112-
                 директор     Ваня    100 руб

             """
            dfs = np.array_split(table, index_to_split)
            dfs = [e for e in dfs if len(e) > 0]

            result_dfs = []
            for df in dfs:
                office = df.iloc[0, :][0]
                df = df.iloc[1:, :]
                df['office'] = office
                result_dfs.append(df)

            result_dfs = [e for e in result_dfs if not e.empty]
            try:
                result_dfs = pd.concat(result_dfs)
                return result_dfs
            except Exception as ex:
                print(ex)
                # print('rogue file---', table)

        index_to_split = get_indexes_to_split(table)

        if not index_to_split:
            return False, table

        splitted_dfs = split_table(table, index_to_split)
        return True, splitted_dfs

    def convert_pdf_to_docx_to_find_info(self, filename: Path) -> Path:
        # переводим пфд в ворд
        assert str(filename).endswith('.pdf'), 'Файл должен быть в PDF !'
        folder = filename.parents[0]

        orig_file_name = filename.name.strip('.pdf')
        new_name = 'temp_to_delete_' + orig_file_name + '.docx'

        pdf2docx.parse(str(filename), str(folder / new_name))
        return folder / new_name

    @staticmethod
    def check_if_columns_ok(cols: tuple) -> bool:
        '''проверяем, есть ли в заголовках таблицы название предприятия и другая инфа'''

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))

        ok_cols = 0
        for col in cols:
            name_salary_position_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|плат[ы, а]|заработная|плата|cреднемесячн[ая, ой]|зарплат[а, ной, ы]|должность)'

            res = re.search(pattern=name_salary_position_pattern, string=col)
            if res:
                ok_cols += 1

        if ok_cols > 1:
            return True

        return False

    def find_ok_cols(self, df: pd.DataFrame) -> pd.DataFrame:

        cols = df.columns

       # если колонки норм, отдаем df
        if self.check_if_columns_ok(cols):
            return df

        i = -1
        for _, row in df.iterrows():
            print(f'-------- ищем в теле {i} ------')
            i += 1
            found_cols = self.check_if_columns_ok(list(row))
            if found_cols:
                print('df at found_cols---------', df)

                df.columns = df.iloc[i, :]
                return df.iloc[i+1:, :]
            if i > 5:
                break

        print('-'*10)
        print('не найдены нормальные колонки')
        print(df.head(3))
        print('-'*10)


#        raise ValueError('Incorrect Headers find_ok_cols --- в df не найдены нормальные колонки')


    def convert_pdf_to_df_and_find_tables(self, filename: str) -> list[pd.DataFrame]:
        #        print('filename in convert_pdf_to_df -----',filename)
        try:
            tables = camelot.read_pdf(str(filename), line_tol=2, joint_tol=10, line_scale=40, copy_text=[
                'v'], pages='1-end')  # , flavor='stream' row_tol=10
            tables = [e.df for e in tables]
            tables = [self.find_ok_cols(e) for e in tables]
            tables = [e for e in tables if type(e) == pd.DataFrame]

            return tables

        except Exception as ex:
            logger.error('file --- %s', filename)
            logger.error('Exception --- %s', ex)

    def detect_headers_in_raw_doc(self, filename) -> list[pd.DataFrame]:

        def get_headers(filename: str) -> list[str]:  # filename:docx

            doc = docx2python(filename)

            table_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|должность)'

            offices = []
            gathering_office_info = ''

            for paragraph in doc.body_runs:  # параграфы в виде вложенных листов

                paragraph = sum(sum(paragraph, []), [])
                paragraph_text = ''
                for e in paragraph:
                    try:
                        paragraph_text += ' ' + e[0] + ' '
                    except IndexError:
                        pass

                paragraph_text = paragraph_text.lower()
                its_table = re.findall(
                    pattern=table_pattern, string=paragraph_text)

                if not its_table:
                    gathering_office_info += paragraph_text

                elif its_table:
                    offices.append(gathering_office_info)
                    gathering_office_info = ''

            return [e for e in offices if e]

        def compile_office_info_and_df(filename, departments) -> list[pd.DataFrame]:
            # все правильно. логика такая, что камелотом лучше парсить!
            # а док только для загов таблиц

            tables = self.convert_pdf_to_df_and_find_tables(filename)

            ok_dfs = []

            print('Количество заголовков --- ', len(departments))
            print('Количество таблиц --- ', len(tables))

            if len(departments) - len(tables) == 1:
                departments.pop()

            if len(departments) == len(tables):
                for table, dep in zip(tables, departments):
                    table['Учреждение'] = dep
                    table['Учреждение'][0] = 'Учреждение'

                    ok_dfs.append(table)

                return ok_dfs

            with open(str(filename) + '.pkl', 'wb') as f:
                pickle.dump('неравные_таблицы_и_учреждения_' +
                            str(filename), f)

        temp_docfile = self.convert_pdf_to_docx_to_find_info(
            filename)  # получили path временного docx файла
        departments = get_headers(temp_docfile)
        dfs = compile_office_info_and_df(filename, departments)
        return dfs

    def parse(self, filename: Path) -> tuple[bool, pd.DataFrame]:
        # пытаемся найти учреждения в теле таблиц

        # TODO: добавить проверку doc или pdf

        tables = self.convert_pdf_to_df_and_find_tables(filename)

        if not tables:
            return False, []

        tables_with_ok_headers = []

        for table in tables:
            res, df = self.table_splitter(table)
            if res:
                tables_with_ok_headers.append(df)

            if not res:
                # идем парсить весь док, чтобы достать учреждения из текста перед таблицей
                dfs = self.detect_headers_in_raw_doc(filename)
                if dfs:
                    for df in dfs:
                        tables_with_ok_headers.append(df)
                break

        # TODO: добавить удаление временного дока

        del tables
        return bool(tables_with_ok_headers), tables_with_ok_headers


class Parser:

    def __init__(self):
        self.cols_we_need = ['name', 'salary', 'position', 'department']
        self.all_docs: list[str]
        self.docx_parser = DocxParser()
        self.pdf_parser = PdfParser()
        self.parse_correct_headers = CorrectHeadersParser()
        self.incorrect_headers_parser = IncorrectHeaders()
        self.data_cleaner = DataCleaner()

    @staticmethod
    def rename_col(col: str) -> str:

        print('col before rename cols --', col)
        col = str(col).lower()
        if re.search(pattern='(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество)', string=col):
            return "name"

        elif re.search(pattern='(рублей|руб|cреднемесячная|зарпл.|плат[ы, а]|заработн[ой, ая] плат[а, ы]|cреднемесячн[ая, ой]|зарплат[а, ной, ы])', string=col):
            return "salary"

        elif re.search(pattern='(должност[ь, и, ей])', string=col):

            return 'position'

        elif re.search(pattern='(предприяти[е,я]|учреждени[е,я]|юридическое лицо)', string=col):
            return 'department'

        return col

    @staticmethod
    def check_if_columns_ok(cols: tuple) -> bool:
        '''проверяем, есть ли в заголовках таблицы название предприятия и другая инфа'''

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))
        ok_cols = 0
        company_found = False
        for col in cols:
            company_pattern = '(предприяти[е,я]|учреждени[е,я]|юридическ[ое,ие])'
            res = re.search(pattern=company_pattern, string=col)
            if res:
                company_found = True
                continue

            name_salary_position_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|плат[ы, а]|заработная|плата|cреднемесячн[ая, ой]|зарплат[а, ной, ы]|должность)'
            res = re.search(pattern=name_salary_position_pattern, string=col)
            if res:
                ok_cols += 1

        if company_found and ok_cols > 1:
            return True
        return False

    def parse_file(self, file: str) -> pd.DataFrame:

        if str(file).endswith('.pdf'):
            tables = self.pdf_parser.convert_pdf_to_df(file)

        elif file.endswith('docx'):
            tables = self.docx_parser.convert_docx_to_df(file)

        else:
            logger.error('Допустимые расширения: pdf, docx')

        parsed_tables = []

        for table in tables:
            # удалить кал из маленьких ячеек?

            columns_ok = self.check_if_columns_ok(table)

            if not columns_ok:
                # колонки непправильные. идем пытаться найти нормальные. для этого мы весь файл передаем в Incorrect и скипаем цикл

                parsing_ok, tables = self.incorrect_headers_parser.parse(
                    Path(file))

                if parsing_ok:
                    for table in tables:
                        # оставляем нужные колонки
                        table.columns = [self.rename_col(
                            col) for col in table.columns]
                        cols_to_leave = [
                            col for col in table.columns if col in self.cols_we_need]
                        cols_to_leave = set(cols_to_leave)
                        table = table[cols_to_leave]
                        # проверяем на наличие вложенных таблиц и фио, разнесенных на несколько стаоблцов
                        table = self.parse_correct_headers.parse(table)
                        # убираем лишние ячейки и символы
                        table = self.data_cleaner.clean_df(table)
                        parsed_tables.append(table)
                    break

                else:
                    with open(f'не_распарсилось_{file}.pkl', 'wb') as f:
                        pickle.dump(file, f)

                    logger.warning('Не удалось распарсить ----', file)

            elif columns_ok:
                # если заголовки ок
                # оставляем только нужные колонки

                table.columns = [self.rename_col(col) for col in table.columns]
                cols_to_leave = [
                    col for col in table.columns if col in self.cols_we_need]
                cols_to_leave = set(cols_to_leave)
                table = table[cols_to_leave]
                # проверяем на наличие вложенных таблиц и фио, разнесенных на несколько стаоблцов
                table = self.parse_correct_headers.parse(table)
                # убираем лишние ячейки и символы
                table = self.data_cleaner.clean_df(table)
                parsed_tables.append(table)

        if isinstance(parsed_tables, list):
            if parsed_tables:
                concat_tables = pd.concat(parsed_tables)
                return concat_tables

        elif isinstance(parsed_tables, pd.DataFrame):
            if not parsed_tables.empty:
                return concat_tables


parser = Parser()

folder = Path('./data_ids/pdf/converted')

for file in os.listdir(folder):
    if file.endswith('.pdf'):  # '100185'
        try:
            res = parser.parse_file(folder / file)
            temp_file = folder / 'cool' / file
            res.to_excel(str(temp_file) + '.xlsx')

        except Exception as ex:
            with open(f'error_{file}.txt', 'w') as f:
                traceback.print_exc(file=f)


# file = r"D:\PROGR\LEARN_PYTHON\Declarator\declarations-parser\data_ids\pdf\converted\100185_2019_Rukovoditeli_podvedomstvennykh_uchrezhdenii_(sport).pdf"
# res = parser.parse_file(file)
# res.to_excel('результат.xlsx')

# with open('result.pkl', 'wb') as f:
#     pickle.dump(res, f)

# for e in os.listdir(folder):
#     if e.endswith('.pdf'):
#         res = parser.parse_file(folder / e)
#         with open(e+'.pkl','wb') as f:
#             pickle.dump(res, f)
