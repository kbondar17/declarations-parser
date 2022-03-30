from docx2python import docx2python
import requests
import re
import time
import os
import urllib.request
from typing import Union
import csv
import statistics
import pdf2docx
import traceback

import pickle

import camelot
import typing
import io
import numpy as np
from docx import Document
import matplotlib
import numpy as np
import matplotlib.pyplot as plt

import pandas as pd
from pathlib import Path
# import tqdm.notebook.tqdm as tqdm
from tqdm import tqdm_notebook
import matplotlib.pyplot as plt
import logging

logging.getLogger('camelot').setLevel('ERROR')

logging.basicConfig(format=u'%(filename)+13s [ LINE:%(lineno)-4s] [ Function %(funcName)s ] :::   %(message)s',
                    level=logging.DEBUG, filename='my_log.log', filemode='w')

my_logger = logging.getLogger(__name__)


class PdfParser:

    @staticmethod
    def convert_pdf_to_df(filename) -> list[pd.DataFrame]:
        tables = camelot.read_pdf(str(filename), line_tol=2, joint_tol=10, line_scale=40, copy_text=[
                                  'v', 'h'], pages='1-end')  # , flavor='stream' row_tol=10
        tables = [e.df for e in tables]
        return tables

    def get_camelot_tables(self, filename):
        tables = camelot.read_pdf(str(filename), line_tol=2, joint_tol=100, line_scale=40, copy_text=[
                                  'v'], pages='1-end')  # , flavor='stream' row_tol=10
        return tables


class CorrectHeadersParser:

    '''класс для парсинга таблиц, у которых на месте колонки, которые нам нужны'''

    def table_splitter(self, table: pd.DataFrame) -> list[pd.DataFrame]:
        '''разделяет таблицы, в которых учреждение указано внутри таблицы'''

        def check_if_same(my_array: list) -> bool:
            '''проверяем одинаковые ли колонки'''

            if len(set(my_array)) > 1:
                return False
            return True

            # first = my_array[0]
            # for e in my_array[1:]:
            #     if e != first:
            #         return False
            # return True

        def get_indexes_to_split(table):
            index_to_split = []
            for e in range(len(table)):
                cols = table.iloc[e, :].values
                if check_if_same(cols):
                    index_to_split.append(e)
            return index_to_split

        def split_table(table: pd.DataFrame, index_to_split: Union[int, list[int]]) -> list[pd.DataFrame]:
            """разделяет таблицу в случае когда название учреждения поместили в середину вот так:

                -должность-  -имя-  -зарплата-
                        -ГБОУ школа 112-
                 директор     Ваня    100 руб

             """
            dfs = np.array_split(table, index_to_split)
            dfs = [e for e in dfs if len(e) > 0]

            result_dfs = []
            for df in dfs:
                office = df.iloc[0, :][0]
                df = df.iloc[1:, :]
                df['office'] = office
                result_dfs.append(df)

            result_dfs = [e for e in result_dfs if not e.empty]
            try:
                result_dfs = pd.concat(result_dfs)
                return result_dfs
            except Exception as ex:
                print(ex)
                print('rogue file---', table)

        index_to_split = get_indexes_to_split(table)

        if not index_to_split:
            return table

        splitted_dfs = split_table(table, index_to_split)
        return splitted_dfs

    def concat_name(self, df: pd.DataFrame) -> pd.DataFrame:
        '''соединяем колонки ФИО, если они в разных'''

        if 'name' not in df.columns:
            return df

        names_df = df['name']

        if isinstance(names_df, str) or isinstance(names_df, pd.Series):
            return df

        # TODO:
        # дропнуть маленькую колонку

        names = [' '.join(e) for e in names_df.values]

        df.drop(columns=['name'], inplace=True)
        df['name'] = names
        return df

    def parse(self, table: pd.DataFrame) -> pd.DataFrame:
        # table = self.concat_name(table)
        table = self.table_splitter(table)
        return table


class DataCleaner:
    """убирает лишние данные"""

    @staticmethod
    def remove_unwanted_symbols(df):
        # TODO: чистка всех колонок
        df = df.applymap(lambda x: str(x).replace('\n', ' '))
        return df

    @staticmethod
    def remove_unwanted_cells(df):
        print('КОЛОНКИ В remove_unwanted_cells ----- ', df.columns)

        # убирает ячейки с нумерацией
        # print('--- DataCleaner.remove_unwanted_cells ---', df.columns)
        # TODO: почему тут только должность?

        # df = df[~df['position'].astype(str).str.isdigit()]
        return df

    @staticmethod
    def remove_short_rows(df: pd.DataFrame):
        # удаляет ряды с недостаточными данными
        # ! должно применяться после выбора норм колонок
        to_remove = []
        for tup in df.itertuples():
            res = [len(str(e)) for e in tup]
            if statistics.mean(res) < 5:
                to_remove.append(tup.Index)
        df.drop(to_remove, inplace=True)

        return df

    def merge_if_three_names(df: pd.DataFrame):
        # TODO: !!!!
        pass

    def clean_df(self, df):
        df = self.remove_unwanted_symbols(df)
        df = self.remove_unwanted_cells(df)

        df = self.remove_short_rows(df)
        # print('!!!',df)

        return df


class DocxParser:

    def get_docx_tables(self, filename, tab_id=None, **kwargs) -> list[pd.DataFrame]:
        """
            filename:   file name of a Word Document
            tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                        When [None] - return a list of DataFrames (parse all tables)
        """
        def read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            return pd.read_csv(vf, **kwargs)

        doc = Document(filename)
        if tab_id is None:
            return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
        else:
            try:
                return read_docx_tab(doc.tables[tab_id], **kwargs)
            except IndexError:
                print(
                    'Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    def convert_docx_to_df(self, filename: str) -> pd.DataFrame:
        assert filename.endswith('docx'), 'Формат должен быть .docx!'

        doc = Document(filename)
        # TODO: тут взять текст, который потом прикрутить к

        doc_tables = self.get_docx_tables(filename)

        return doc_tables


class IncorrectHeaders:
    """класс для таблиц с неопределенными заголовками.
        1. Пытаемся найти название учреждений в объединенных ячейках.
        2. Если не получается, для учреждения берем текст, предшествующий таблице. 
    """

    # если прошелся по таблице и нашел вложения внутри - пусть это будет офис.
    # если не нашел - берем название офиса из абзацев вокруг таблиц (если их число плюс минус совпадает)

    def __init__(self):
        pass
        # TODO: добавить обработку чисто docx

        # self.docx_parser = DocxParser()
        # self.pdf_parser = PdfParser()

    @staticmethod
    def drop_col_with_N(df: pd.DataFrame):
        expr = '(№|п/п)'
        for c in df.columns:

            if re.search(expr, str(c)):
                df.drop(columns=c, inplace=True)
        return df

    @staticmethod
    def drop_short_cols(df: pd.DataFrame):
        len_df = df.applymap(len)
        df.to_excel('drop_test.xlsx')
        for c in df.columns:
            if len_df[c].mean() < 4:
                df.drop(columns=c, inplace=True)
        return df

    @staticmethod
    def drop_short_headers(df: pd.DataFrame) -> pd.DataFrame:
        for i in range(3):
            cols = list(map(str, df.columns))
            cols = list(map(len, cols))
            if statistics.mean(cols) < 3 and i < 2:
                df.columns = df.iloc[0, :]
                # df = df.iloc[1:, :]
            else:
                return df

    def table_splitter(self, table: pd.DataFrame) -> tuple[bool, list[pd.DataFrame]]:
        '''разделяет таблицы, в которых учреждение указано внутри таблицы'''

        def check_if_same(my_array: list) -> bool:
            '''проверяем одинаковые ли колонки'''

            if len(set(my_array)) > 1:
                return False
            return True

        def get_indexes_to_split(table):
            '''определяем индекс строки таблицы, по которому надо разделить'''

            index_to_split = []
            for e in range(len(table)):
                cols = table.iloc[e, :].values
                if check_if_same(cols):
                    index_to_split.append(e)
            return index_to_split

            # если разделили и нашли офис - True
        def split_table(table: pd.DataFrame, index_to_split: Union[int, list[int]]) -> tuple[bool, list[pd.DataFrame]]:
            """разделяет таблицу в случае когда название учреждения поместили в середину вот так:
                -должность-  -имя-  -зарплата-
                        -ГБОУ школа 112-
                 директор     Ваня    100 руб

             """
            dfs = np.array_split(table, index_to_split)
            dfs = [e for e in dfs if len(e) > 0]

            result_dfs = []
            for df in dfs:
                office = df.iloc[0, :][0]
                df = df.iloc[1:, :]
                df['office'] = office
                result_dfs.append(df)

            result_dfs = [e for e in result_dfs if not e.empty]
            try:
                result_dfs = pd.concat(result_dfs)
                return result_dfs
            except Exception as ex:
                print(ex)
                # print('rogue file---', table)

        index_to_split = get_indexes_to_split(table)

        if not index_to_split:
            return False, table

        splitted_dfs = split_table(table, index_to_split)
        return True, splitted_dfs

    def convert_pdf_to_docx_to_find_info(self, filename: Path) -> Path:
        # переводим пфд в ворд
        assert str(filename).endswith('.pdf'), 'Файл должен быть в PDF !'
        folder = filename.parents[0]

        orig_file_name = filename.name.strip('.pdf')
        new_name = 'temp_to_delete_' + orig_file_name + '.docx'

        pdf2docx.parse(str(filename), str(folder / new_name))
        return folder / new_name

    @staticmethod
    def check_if_columns_ok(cols: tuple) -> bool:
        '''проверяем, есть ли в заголовках таблицы нужная инфа'''

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))

        ok_cols = 0
        for col in cols:
            name_salary_position_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|плат[ы, а]|заработная|плата|cреднемесячн[ая, ой]|зарплат[а, ной, ы]|должность)'

            res = re.search(pattern=name_salary_position_pattern, string=col)
            if res:
                ok_cols += 1

        if ok_cols > 1:
            return True

        return False

    def find_ok_cols(self, df: pd.DataFrame) -> dict['df':pd.DataFrame, 'if_ok_cols':bool]:

        cols = df.columns
       # если колонки норм, отдаем df
        if self.check_if_columns_ok(cols):
            return {'df': df, 'if_ok_cols': True}

        i = -1
        for _, row in df.iterrows():
            i += 1
            found_cols = self.check_if_columns_ok(list(row))
            if found_cols:
                df.columns = df.iloc[i, :]
                return {'df': df.iloc[i+1:, :], 'if_ok_cols': True}

            if i > 5:
                break

        # если не ок
        return {'df': df, 'if_ok_cols': False}

    @staticmethod
    def if_office_in_cols(dfs: list[pd.DataFrame]) -> bool:
        for df in dfs:
            cols = df.columns

            cols = list(map(str, cols))
            cols = list(map(str.lower, cols))

            office_pattern = '(предприяти[е,я]|учреждени[е,я]|юридическ[ие, ое])'

            if not any([re.search(pattern=office_pattern, string=col) for col in cols]):
                return False

        return True

    def convert_pdf_to_dfs(self, filename: str) -> list[pd.DataFrame]:

        try:
            tables = camelot.read_pdf(str(filename), line_tol=2, joint_tol=10, line_scale=40, copy_text=[
                'v', 'h'], pages='1-end')  # , flavor='stream' row_tol=10
            tables = [e.df for e in tables]
            return tables

        except Exception as ex:
            my_logger.error('file --- %s', filename)
            my_logger.error('Exception --- %s', ex)

    def detect_headers_in_raw_doc(self, filename) -> list[pd.DataFrame]:

        def get_headers(filename: str) -> list[str]:  # filename:docx

            doc = docx2python(filename)

            table_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|должность)'

            offices = []
            gathering_office_info = ''

            for paragraph in doc.body_runs:  # параграфы в виде вложенных листов

                paragraph = sum(sum(paragraph, []), [])
                paragraph_text = ''
                for e in paragraph:
                    try:
                        paragraph_text += ' ' + e[0] + ' '
                    except IndexError:
                        pass

                paragraph_text = paragraph_text.lower()
                its_table = re.findall(
                    pattern=table_pattern, string=paragraph_text)

                if not its_table:
                    gathering_office_info += paragraph_text

                elif its_table:
                    offices.append(gathering_office_info)
                    gathering_office_info = ''

            return [e for e in offices if e]

        def compile_office_info_and_df(filename: Path, departments: list) -> typing.Union[None, list[pd.DataFrame]]:
            # все правильно. логика такая, что камелотом лучше парсить!
            # а док только для загов таблиц

            tables = self.convert_pdf_to_dfs(filename)

            ok_dfs = []

            print('Количество заголовков --- ', len(departments))
            print('Количество таблиц --- ', len(tables))

            if len(departments) - len(tables) == 1:
                departments.pop()

            if len(departments) == len(tables):
                for table, dep in zip(tables, departments):
                    table['Учреждение'] = dep
                    table['Учреждение'][0] = 'Учреждение'

                    ok_dfs.append(table)

                return ok_dfs

            with open(str(filename) + '.txt', 'r') as f:
                text = f'Разное число таблиц ({len(tables)}) и учреждений ({len(departments)})'
                f.write(text)

#            raise ValueError(f'Разное число таблиц ({len(tables)}) и учреждений ({len(departments)})')

        temp_docfile = self.convert_pdf_to_docx_to_find_info(
            filename)  # получили path временного docx файла
        departments = get_headers(temp_docfile)
        dfs = compile_office_info_and_df(filename, departments)
        return dfs

    def concatenate_if_possible(self, dfs: list[dict['df':pd.DataFrame, 'if_ok_cols':bool]]) -> list[pd.DataFrame]:

        all_oks = [e['if_ok_cols'] for e in dfs]

        if all(all_oks):
            return [e['df'] for e in dfs]

        result_df = []
        df_to_concat = pd.DataFrame()
        for df_info in dfs:
            if df_info['if_ok_cols']:
                if not df_to_concat.empty:
                    result_df.append(df_to_concat)
                df_to_concat = df_info['df']

            # оставляем только таблицы, у которых совпадает число колонок
            # с df у которых мы колонки нашли
            # если не нашли колонки и не к чему присоединять - дропаем

            elif not df_info['if_ok_cols'] and not df_to_concat.empty \
                    and len(df_to_concat.columns) == len(df_info['df'].columns):
                df_info['df'].columns = df_to_concat.columns
                df_to_concat = pd.concat([df_to_concat, df_info['df']])

        result_df.append(df_to_concat)

        return result_df

    # Path - относительный
    def parse(self, filename: Path) -> tuple[bool, pd.DataFrame]:
        # пытаемся найти учреждения в теле таблиц

        # TODO: добавить проверку doc или pdf

        # должны быть просто таблицы
        # и вся обработка должна быть тут, по этапам. иначе макароны
        tables = self.convert_pdf_to_dfs(filename)
        # дропаем маленькие колонки
        tables = [self.drop_short_headers(e) for e in tables]
        tables = [self.drop_col_with_N(e) for e in tables]
        tables = [e for e in tables if type(e) == pd.DataFrame]
        tables = [self.drop_short_cols(e) for e in tables]

        """
        Есть все таблицы. У некоторых нет вообще заголовков. 
        Присоединяем их к тем у кого есть заголовки.
        Получаем таблицы.
        Если у них нет учреждений - идем парсить в док. 
        
        """

        # у нас тут лист словарей {df:bool}. к каждой таблице мы должны приделать True или False
        tables = [self.find_ok_cols(e) for e in tables]
        # TODO: если нет ни одной таблицы с ок загами -> скипаем все

        at_least_one_table_ok = any([e['if_ok_cols'] for e in tables])
        if not at_least_one_table_ok:
            return False, []

        # теперь надо склеить таблицы, если есть таблицы с ок колонками
        tables = self.concatenate_if_possible(tables)

        # проверяем есть ли учреждение
        if self.if_office_in_cols(tables):
            return True, tables

        # если нет - парсим док.

        tables_with_ok_headers = []

        for table in tables:
            res, df = self.table_splitter(table)
            if res:
                tables_with_ok_headers.append(df)

            if not res:
                # идем парсить весь док, чтобы достать учреждения из текста перед таблицей
                dfs = self.detect_headers_in_raw_doc(filename)
                if not dfs:
                    return False, []
                for df in dfs:
                    tables_with_ok_headers.append(df)
                break

        # TODO: переделать удаление временного дока
        return bool(tables_with_ok_headers), tables_with_ok_headers


class Parser:

    def __init__(self):
        self.cols_we_need = ['name', 'salary', 'position', 'department']
        self.all_docs: list[str]
        self.docx_parser = DocxParser()
        self.pdf_parser = PdfParser()
        self.parse_correct_headers = CorrectHeadersParser()
        self.incorrect_headers_parser = IncorrectHeaders()
        self.data_cleaner = DataCleaner()

    @staticmethod
    def rename_col(col: str) -> str:

        print('col before rename cols --', col)
        col = str(col).lower()
        if re.search(pattern='(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество)', string=col):
            return "name"

        elif re.search(pattern='(рублей|руб|cреднемесячная|зарпл.|плат[ы, а]|заработн[ой, ая] плат[а, ы]|cреднемесячн[ая, ой]|зарплат[а, ной, ы])', string=col):
            return "salary"

        elif re.search(pattern='(должност[ь, и, ей])', string=col):

            return 'position'

        elif re.search(pattern='(предприяти[е,я]|учреждени[е,я]|юридическое лицо)', string=col):
            return 'department'

        return col

    @staticmethod
    def check_if_columns_ok(cols: tuple) -> bool:
        '''проверяем, есть ли в заголовках таблицы название предприятия и другая инфа'''

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))
        ok_cols = 0
        company_found = False
        for col in cols:
            company_pattern = '(предприяти[е,я]|учреждени[е,я]|юридическ[ое,ие])'
            res = re.search(pattern=company_pattern, string=col)
            if res:
                company_found = True
                continue

            name_salary_position_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|плат[ы, а]|заработная|плата|cреднемесячн[ая, ой]|зарплат[а, ной, ы]|должность)'
            res = re.search(pattern=name_salary_position_pattern, string=col)
            if res:
                ok_cols += 1

        if company_found and ok_cols > 1:
            return True
        return False

    def parse_file(self, file: str) -> pd.DataFrame:

        if str(file).endswith('.pdf'):
            tables = self.pdf_parser.convert_pdf_to_df(file)

        elif file.endswith('docx'):
            tables = self.docx_parser.convert_docx_to_df(file)

        else:
            my_logger.error('Допустимые расширения: pdf, docx')

        parsed_tables = []

        for table in tables:
            # удалить кал из маленьких ячеек?

            columns_ok = self.check_if_columns_ok(table)

            if not columns_ok:
                # колонки непправильные. идем пытаться найти нормальные. для этого мы весь файл передаем в Incorrect и скипаем цикл

                parsing_ok, tables = self.incorrect_headers_parser.parse(
                    Path(file))

                if parsing_ok:
                    for table in tables:
                        table.reset_index(inplace=True)
                        # оставляем нужные колонки
                        # продолжить дебажить
                        table.columns = [self.rename_col(
                            col) for col in table.columns]
                        cols_to_leave = [
                            col for col in table.columns if col in self.cols_we_need]
                        cols_to_leave = set(cols_to_leave)
                        table = table[cols_to_leave]
                        # проверяем на наличие вложенных таблиц и фио, разнесенных на несколько стаоблцов
                        table = self.parse_correct_headers.parse(table)
                        # убираем лишние ячейки и символы
                        table = self.data_cleaner.clean_df(table)
                        parsed_tables.append(table)

                    break

                else:
                    # TODO: сохранить файл в папку нераспаршенных
                    my_logger.warning('Не удалось распарсить ----', file)

            elif columns_ok:
                # если заголовки ок
                # оставляем только нужные колонки
                table.reset_index(inplace=True)
                table.columns = [self.rename_col(col) for col in table.columns]
                cols_to_leave = [
                    col for col in table.columns if col in self.cols_we_need]
                cols_to_leave = set(cols_to_leave)
                table = table[cols_to_leave]
                # проверяем на наличие вложенных таблиц и фио, разнесенных на несколько стаоблцов
                table = self.parse_correct_headers.parse(table)
                # убираем лишние ячейки и символы
                table = self.data_cleaner.clean_df(table)
                parsed_tables.append(table)

        if isinstance(parsed_tables, list):
            if parsed_tables:
                concat_tables = pd.concat(parsed_tables)
                return concat_tables

        elif isinstance(parsed_tables, pd.DataFrame):
            if not parsed_tables.empty:
                return concat_tables


parser = Parser()
folder = Path('./data_ids/pdf/converted')
# with open('')


# file = r"D:\PROGR\LEARN_PYTHON\Declarator\declarations-parser\data_ids\pdf\converted\83301_2018_Rukovoditeli,_zamestiteli_i_glavnye_bukhgaltery_podvedomstvennykh_uchrezhdenii_(FGBU).pdf"
file = r"D:\PROGR\LEARN_PYTHON\Declarator\declarations-parser\data_ids\pdf\converted\83300_2018_Rukovoditeli,_zamestiteli_i_glavnye_bukhgaltery_podvedomstvennykh_uchrezhdenii_(FGUP).pdf"
file = r"D:\PROGR\LEARN_PYTHON\Declarator\declarations-parser\data_ids\pdf\converted\83333_2018_Rukovoditeli,_zamestiteli_i_glavnye_bukhgaltery_podvedomstvennykh_uchrezhdenii.pdf"
file = r"D:\PROGR\LEARN_PYTHON\Declarator\declarations-parser\data_ids\pdf\83334_2018_Rukovoditeli,_zamestiteli_i_glavnye_bukhgaltery_podvedomstvennykh_uchrezhdenii.pdf"

res = parser.parse_file(file)
res.to_excel('big_test.xlsx')

# for file in os.listdir(folder):
#     if file.endswith('.pdf'):  # '100185'
#         try:
#             res = parser.parse_file(folder / file)
#             temp_file = folder / 'cool' / file
#             res.to_excel(str(temp_file) + '.xlsx')

#         except Exception as ex:
#             with open(f'error_{file}.txt', 'w') as f:
#                 traceback.print_exc(file=f)