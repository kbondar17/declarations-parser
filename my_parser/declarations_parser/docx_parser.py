
import csv
import io
import re
import statistics
import typing
from typing import Union

import pandas as pd
from docx import Document
from docx2python import docx2python
from my_parser.declarations_parser.config import get_logger

logger = get_logger(__name__)


class HeaderFinder:

    def find_and_split_departments(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        отдает df с учреждением, если оно было записано в середине таблицы
        """

        def _add_department_info_to_df(df: pd.DataFrame, dep_info: list[dict[int, str]]) -> pd.DataFrame:
            df['department'] = None
            for data in dep_info:
                df.at[data['index'], 'department'] = data['dep']
            df = df.dropna(axis=1, how='all').fillna(method='ffill', axis=0)
            return df

        # [index, department]
        def _find_department_in_table(df: pd.DataFrame) -> list[dict[int, str]]:
            departments_n_indexes = []
            for row in df.itertuples():
                index = row[0]
                row = list(row)[1:-1]
                row = [e for e in row if len(str(e)) > 4]
                if len(set(row)) < 2:
                    if not all([type(e) in [int, float] for e in row]):
                        if statistics.mean([len(e) for e in row]) > 4:
                            departments_n_indexes.append(
                                {'index': index, 'dep': row[0]})
            return departments_n_indexes

        dep_info = _find_department_in_table(df)

        if not dep_info:
            return df

        df = _add_department_info_to_df(df, dep_info)
        return df


class DocxParser:
    def __init__(self) -> None:
        self.header_finder = HeaderFinder()

    def get_docx_tables(self, filename, tab_id=None, **kwargs) -> list[pd.DataFrame]:
        """
            filename:   file name of a Word x Document
            tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                        When [None] - return a list of DataFrames (parse all tables)
        """
        def read_docx_tab(tab, **kwargs):
            vf = io.StringIO()
            writer = csv.writer(vf)
            for row in tab.rows:
                writer.writerow(cell.text for cell in row.cells)
            vf.seek(0)
            return pd.read_csv(vf, header=None, **kwargs)

        doc = Document(filename)
        if tab_id is None:
            return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
        else:
            try:
                return read_docx_tab(doc.tables[tab_id], **kwargs)
            except IndexError:
                print(
                    'Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                raise

    def convert_docx_to_df(self, filename: str) -> pd.DataFrame:
        assert filename.endswith('docx'), 'Формат должен быть .docx!'

        doc_tables = self.get_docx_tables(filename)

        return doc_tables

    @staticmethod
    def find_ok_cols(cols: tuple) -> Union[dict[int, str], None]:
        """
        ищет в теле df наиболее вероятный ряд с заголовками
        если находит - отдает словарь {номер_столбца:новое название}
        """

        cols = list(map(str, cols))
        cols = list(map(str.lower, cols))

        if len(set(cols)) == 1:
            return

        ok_cols = 0
        result = {}  # 'n_of_col':'new_name'
        for i, col in enumerate(cols):
            if len(result) > 5 or i > 20:
                return result

            if re.search(pattern='(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество)', string=col) and 'name' not in result.values():
                result[i] = 'name'
                ok_cols += 1

            elif re.search(pattern='(рублей|руб|cреднемесячная|зарпл.|плат[ы, е, а]|заработн[ой, ая] плат[а, ы]|cреднемесячн[ая, ой]|зарплат[а, ной, ы])', string=col):
                result[i] = "salary"
                ok_cols += 1

            elif re.search(pattern='(должност[ь, и, ей])', string=col) and 'position' not in result.values():
                result[i] = 'position'
                ok_cols += 1

            elif re.search(pattern='(предприяти[е,я]|учреждени[е,я]|юридическ|организаци|наименование [оу, мо])', string=col) and 'department' not in result.values():
                result[i] = 'department'
                ok_cols += 1

        if len(result) >= 2:
            return result
        return False

    def detect_headers(self, df: pd.DataFrame) -> dict[bool, pd.DataFrame]:
        """
        ищем заголовки в теле df. если находим - переименовываем колонки.
        отдаем словарь {есть ли в df норм заголовки : df} 
        """

        for row in df.itertuples():
            index = row[0]
            new_cols = self.find_ok_cols(row[1:])
            if new_cols:
                break

        if not new_cols:
            return {'ok_cols': False, 'df': df}

        for k, v in new_cols.items():
            df.rename(columns={df.columns[k]: v}, inplace=True)

        df = df.iloc[index:, :]
        return {'ok_cols': True, 'df': df}

    def concatenate_if_possible(self, dfs: list[dict['df':pd.DataFrame, 'ok_cols':bool]]) -> list[pd.DataFrame]:
        """
        может статься так, что на первой страницы у таблицы есть заголовки, а на второй уже нет.
        конкатинируем их в таком случае.
        """

        all_oks = [e['ok_cols'] for e in dfs]
        if all(all_oks):
            return [e['df'] for e in dfs]

        result_df = []
        df_to_concat = pd.DataFrame()
        for df_info in dfs:
            if df_info['ok_cols']:
                if not df_to_concat.empty:
                    result_df.append(df_to_concat)
                df_to_concat = df_info['df']

            elif not df_info['ok_cols'] and not df_to_concat.empty \
                    and len(df_to_concat.columns) == len(df_info['df'].columns):
                df_info['df'].columns = df_to_concat.columns
                df_to_concat = pd.concat([df_to_concat, df_info['df']])

        result_df.append(df_to_concat)

        return result_df

    @staticmethod
    def get_offices_from_doc(filename: str) -> list[str]:
        """
        если в таблице нет колонки с учреждением, берем текст перед таблицей.
        """
        doc = docx2python(filename)
        table_pattern = '(фамилия|имя|фио|ф\.и\.о\.|ф\.и\.о|отчество|должность)'
        offices = []
        gathering_office_info = ''

        for paragraph in doc.body_runs:
            # идем по параграфам. как доходим до таблицы - берем текст перед ней
            paragraph = sum(sum(paragraph, []), [])
            paragraph_text = ''
            for e in paragraph:
                try:
                    paragraph_text += ' ' + e[0] + ' '
                except IndexError:
                    pass

            paragraph_text = paragraph_text.lower()
            its_table = re.findall(
                pattern=table_pattern, string=paragraph_text)

            if not its_table:
                gathering_office_info += paragraph_text

            elif its_table:
                offices.append(gathering_office_info)
                gathering_office_info = ''

        return [e for e in offices if e]

    @staticmethod
    def compile_office_info_and_df(departments: list, tables: list[pd.DataFrame]) -> typing.Union[None, list[pd.DataFrame]]:
        """Соединяем таблицы и куски текста перед ними"""

        ok_dfs = []

        if len(departments) - len(tables) == 1:
            departments.pop()

        if len(departments) == len(tables):
            for table, dep in zip(tables, departments):
                table['department'] = dep
                table['department'][0] = 'Учреждение'
                ok_dfs.append(table)

            return ok_dfs

        logger.warning(
            f'Разное число таблиц ({len(tables)}) и учреждений ({len(departments)})')

        raise ValueError(
            f'Разное число таблиц ({len(tables)}) и учреждений ({len(departments)})')

    def parse_file(self, filepath: str) -> list[pd.DataFrame]:
        dfs = self.get_docx_tables(filepath)
        dfs = [self.detect_headers(df) for df in dfs]
        dfs = self.concatenate_if_possible(dfs)

        how_many_dfs_with_departments = sum(
            [1 for e in dfs if 'department' in e.columns])
        # если учреждений больше, чем у половины таблиц, то отдаем
        if how_many_dfs_with_departments > len(dfs) // 2:
            return dfs

        else:
            dfs = [self.header_finder.find_and_split_departments(
                df) for df in dfs]

            how_many_dfs_with_departments = sum(
                [1 for e in dfs if 'department' in e.columns])
            # если учреждений больше, чем у половины таблиц, то отдаем
            if how_many_dfs_with_departments > len(dfs) // 2:
                return dfs

            departments = self.get_offices_from_doc(filepath)
            dfs = self.compile_office_info_and_df(
                tables=dfs, departments=departments)

            return dfs
